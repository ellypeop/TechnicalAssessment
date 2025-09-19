from __future__ import annotations

from pathlib import Path
from typing import List

from docx import Document


def extract_docx_text(path: Path) -> str:
	if not path.exists() or not path.is_file():
		return ""
	doc = Document(str(path))
	parts: List[str] = []
	for p in doc.paragraphs:
		text = (p.text or "").strip()
		if text:
			parts.append(text)
	# Extract basic table text as lines
	for table in doc.tables:
		for row in table.rows:
			cells = [c.text.strip() for c in row.cells]
			line = " | ".join([c for c in cells if c])
			if line:
				parts.append(line)
	return "\n".join(parts).strip()
