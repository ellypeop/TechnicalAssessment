from __future__ import annotations

from pathlib import Path
from typing import List, Dict
import re

from docx import Document

from .schema import SupplierResponseSchema

# Core noun signals common to RFP "asks"
KEY_TERMS = [
    "requirement", "requirements", "evaluation", "criteria", "question", "questions",
    "scope", "deliverable", "service levels", "sla", "kpi", "technical", "commercial",
    "compliance", "licensing", "license", "cost", "pricing", "roi", "implementation",
    "timeline", "support", "maintenance", "integration", "api", "security", "encryption",
    "retention", "gdpr", "data protection", "case study", "reference"
]

# Lines that imply supplier action (keep if any verb matches)
ACTION_VERBS = [
    "must", "required", "require", "shall", "should", "provide", "submit", "include",
    "demonstrate", "integrate", "support", "comply", "detail", "explain", "present",
    "show", "describe", "outline", "evidence", "evidenced"
]

HEADING_PATTERNS = [
    r"^\s*(section|part|chapter|appendix)\s+\d+\b",
    r"^\s*\d+(\.\d+)*\s+[A-Z][A-Za-z]",     # numbered headings
    r"^\s*[A-Z][A-Z\s]{4,}$",               # ALL CAPS headings
    # Title-Case short headings without sentence punctuation (up to ~10 words)
    r"^(?:[A-Z][a-z]+(?:\s|$)){1,10}$"
]

DATE_PATTERNS = [
    # Ordinals supported (e.g., "16th September 2025")
    r"\b\d{1,2}(st|nd|rd|th)?\s+\w+\s+\d{4}\b",
    r"\b\d{4}-\d{2}-\d{2}\b",
    r"\b\d{2}/\d{2}/\d{4}\b",
    # Agenda/timing words often used as non-asks
    r"\b(valid (from|until)|deadline|submission date|presentation|q&a|q & a|timeline)\b"
]

BOILERPLATE_TERMS = [
    "confidential", "copyright", "all rights reserved", "table of contents",
    "document control", "revision history", "company registered", "registered office"
]


def _looks_like_heading(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    return any(re.search(p, t, flags=re.IGNORECASE) for p in HEADING_PATTERNS)


def _looks_like_date_or_boilerplate(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False
    low = t.lower()
    if any(re.search(p, t, flags=re.IGNORECASE) for p in DATE_PATTERNS):
        return True
    return any(term in low for term in BOILERPLATE_TERMS)


def _is_candidate_line(text: str) -> bool:
    """
    Decide if a line is a potential supplier requirement.
    Reject headings/dates/boilerplate; prefer lines with action verbs or strong nouns.
    """
    t = (text or "").strip()
    if not t:
        return False
    if _looks_like_heading(t) or _looks_like_date_or_boilerplate(t):
        return False

    low = t.lower()

    # Keep only if it contains action verbs OR core terms
    has_action = any(re.search(rf"\b{v}\b", low) for v in ACTION_VERBS)
    has_key = any(k in low for k in KEY_TERMS)

    # Bullets allowed only if they ALSO have action/key signal
    is_bullet = bool(re.match(r"^\s*(\d+\.|[-\*\u2022•])\s+", t))

    return has_action or (has_key and (is_bullet or len(t.split()) > 4))


def _preprocess_text(text: str) -> str:
    # Remove obvious headings/dates/boilerplate and collapse whitespace
    if _looks_like_heading(text) or _looks_like_date_or_boilerplate(text):
        return ""
    text = re.sub(r"\s+", " ", text or "")
    return text.strip()


def extract_criteria_from_docx(path: Path, max_items: int = 20) -> List[str]:
    """
    Extract actionable supplier requirements from a .docx RFP.
    Special handling: force-capture the checklist under
    'Your presentation should cover the following' (these are explicit asks).
    """
    doc = Document(str(path))
    lines: List[str] = []

    # Flag to capture bullets immediately following the checklist intro
    capture_presentation_list = False

    for p in doc.paragraphs:
        raw = p.text or ""
        # Trigger checklist capture if we see the intro line verbatim (case-insensitive)
        if "your presentation should cover the following" in raw.lower():
            capture_presentation_list = True
            continue

        if capture_presentation_list:
            # Stop when we hit a blank or a new heading-like line
            if not raw.strip() or _looks_like_heading(raw):
                capture_presentation_list = False
            else:
                tt = _preprocess_text(raw)
                if tt:
                    # These bullets are genuine asks; keep even if verbs are absent
                    lines.append(tt)
                continue

        # Normal path: preprocess and apply candidate filter
        t = _preprocess_text(raw)
        if t and _is_candidate_line(t):
            lines.append(t)

    # Also scan tables for asks
    for table in doc.tables:
        for row in table.rows:
            for c in row.cells:
                cell_text = _preprocess_text(c.text or "")
                if cell_text and _is_candidate_line(cell_text):
                    lines.append(cell_text)

    # De-duplicate (case-insensitive), preserve order
    seen = set()
    unique: List[str] = []
    for ln in lines:
        k = ln.lower()
        if k in seen:
            continue
        seen.add(k)
        unique.append(ln)

    return unique[:max_items]


# Broader, RFP-aligned categorisation to bucket the asks for downstream prompts
CATEGORY_KEYWORDS: Dict[str, List[str]] = {
    "capabilities": [
        "feature", "function", "capability", "workflow", "report",
        "ai", "first-draft", "library", "web-based"
    ],
    "integrations": [
        "integration", "api", "sso", "single sign-on", "erp", "crm",
        "microsoft teams", "salesforce", "shared drives"
    ],
    "users_roles": [
        "user", "role", "permission", "access", "admin", "super user", "reviewer"
    ],
    "security": [
        "security", "gdpr", "uk gdpr", "iso", "encryption",
        "data protection", "backup", "retention", "privacy"
    ],
    "commercials": [
        "price", "pricing", "cost", "license", "licensing",
        "contract", "sla", "roi", "return on investment", "cost breakdown"
    ],
    "delivery": [
        "implementation", "timeline", "migration", "training", "onboarding",
        "plan", "approach", "schedule"
    ],
    "support": [
        "support", "helpdesk", "ticket", "response", "uptime", "maintenance",
        "service", "post-implementation"
    ],
    "evidence": [
        "case study", "case studies", "reference", "references", "evidence", "proof", "example"
    ],
    "submission": [
        "submission", "format", "appendix", "document", "response",
        "confirm attendance", "presentation", "q&a", "q & a", "deadline"
    ],
    "partnership": [
        "partner", "partnership", "collaborat", "account management", "governance",
        "roadmap", "scalability", "long-term"
    ],
}


def categorize_to_schema(items: List[str]) -> SupplierResponseSchema:
    schema = SupplierResponseSchema()
    for it in items:
        low = it.lower()
        assigned = False
        for field, kws in CATEGORY_KEYWORDS.items():
            if any(kw in low for kw in kws):
                getattr(schema, field).append(it)
                assigned = True
                break
        if not assigned:
            # default bucket
            schema.capabilities.append(it)
    return schema
